"""Utilities for exporting analysis results to various formats."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Optional
from datetime import datetime

import pandas as pd
import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")


def export_statistics_to_word(
    output_path: str | Path,
    overall_stats: Optional[Dict[str, Any]] = None,
    slope_stats: Optional[pd.DataFrame] = None,
    landcover_stats: Optional[pd.DataFrame] = None,
    height_stats: Optional[pd.DataFrame] = None,
    shadow_stats: Optional[Dict[str, Any]] = None,
    additional_stats: Optional[Dict[str, pd.DataFrame]] = None,
    images_dir: Optional[str | Path] = None,
    include_images: bool = True,
    title: str = "SAOCOM DEM Validation Analysis Results"
) -> None:
    """
    Export comprehensive analysis statistics to a Word document.

    Parameters
    ----------
    output_path : str | Path
        Path where the Word document will be saved
    overall_stats : dict, optional
        Overall validation statistics (bias, RMSE, NMAD) for reference DEMs
        Expected format: {
            'TINItaly': {'bias': x, 'rmse': y, 'nmad': z, 'n_points': n},
            'Copernicus': {'bias': x, 'rmse': y, 'nmad': z, 'n_points': n}
        }
    slope_stats : pd.DataFrame, optional
        Statistics stratified by slope category
    landcover_stats : pd.DataFrame, optional
        Statistics stratified by land cover type
    height_stats : pd.DataFrame, optional
        Statistics stratified by elevation range
    shadow_stats : dict, optional
        Shadow and geometry statistics by category
    additional_stats : dict, optional
        Additional statistics tables to include (name: dataframe pairs)
    images_dir : str | Path, optional
        Directory containing result images to embed
    include_images : bool, default=True
        Whether to include images in the document
    title : str, default="SAOCOM DEM Validation Analysis Results"
        Document title

    Raises
    ------
    ImportError
        If python-docx is not installed
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for Word export. "
            "Install with: pip install python-docx"
        )

    # Create new document
    doc = Document()

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")  # Blank line

    # Section 1: Overall Validation Statistics
    if overall_stats:
        doc.add_heading("1. Overall Validation Statistics", level=1)
        doc.add_paragraph(
            "Comparison of SAOCOM heights against reference DEMs. "
            "Metrics computed after calibration and outlier removal."
        )

        # Create table for overall statistics
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Reference DEM', 'N Points', 'Bias (m)', 'RMSE (m)', 'NMAD (m)']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            _format_cell(header_cells[i], bold=True)

        # Data rows
        for ref_name, stats in overall_stats.items():
            row_cells = table.add_row().cells
            row_cells[0].text = ref_name
            row_cells[1].text = f"{stats.get('n_points', 0):,}"
            row_cells[2].text = f"{stats.get('bias', 0):+.2f}"
            row_cells[3].text = f"{stats.get('rmse', 0):.2f}"
            row_cells[4].text = f"{stats.get('nmad', 0):.2f}"

        doc.add_paragraph("")  # Blank line

        # Add interpretation note
        note = doc.add_paragraph()
        note.add_run("Note: ").bold = True
        note.add_run(
            "NMAD (Normalized Median Absolute Deviation) is a robust accuracy "
            "metric less sensitive to outliers than RMSE."
        )
        doc.add_paragraph("")

    # Section 2: Slope Stratification
    if slope_stats is not None and not slope_stats.empty:
        doc.add_heading("2. Accuracy by Slope Category", level=1)
        doc.add_paragraph(
            "Analysis of height accuracy stratified by terrain slope. "
            "Steeper slopes typically show larger residuals due to "
            "geometric distortions and shadow effects."
        )

        _add_dataframe_table(doc, slope_stats, "Slope Category")
        doc.add_paragraph("")

    # Section 3: Land Cover Stratification
    if landcover_stats is not None and not landcover_stats.empty:
        doc.add_heading("3. Accuracy by Land Cover Type", level=1)
        doc.add_paragraph(
            "Analysis of height accuracy stratified by land cover classification. "
            "Different surface types exhibit varying radar backscatter properties."
        )

        _add_dataframe_table(doc, landcover_stats, "Land Cover Type")
        doc.add_paragraph("")

    # Section 4: Elevation Stratification
    if height_stats is not None and not height_stats.empty:
        doc.add_heading("4. Accuracy by Elevation Range", level=1)
        doc.add_paragraph(
            "Analysis of height accuracy stratified by elevation categories."
        )

        _add_dataframe_table(doc, height_stats, "Elevation Range")
        doc.add_paragraph("")

    # Section 5: Shadow/Geometry Statistics
    if shadow_stats:
        doc.add_heading("5. Radar Geometry and Shadow Statistics", level=1)
        doc.add_paragraph(
            "Statistics for different radar viewing geometry categories, "
            "including shadow and layover zones."
        )

        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        headers = ['Category', 'Count', 'Bias (m)', 'RMSE (m)', 'NMAD (m)']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            _format_cell(header_cells[i], bold=True)

        # Data rows
        for category, stats in shadow_stats.items():
            if stats.get('count', 0) > 0:
                row_cells = table.add_row().cells
                row_cells[0].text = str(category)
                row_cells[1].text = f"{stats['count']:,}"
                row_cells[2].text = f"{stats.get('bias', 0):+.2f}"
                row_cells[3].text = f"{stats.get('rmse', 0):.2f}"
                row_cells[4].text = f"{stats.get('nmad', 0):.2f}"

        doc.add_paragraph("")

    # Section 6: Additional Statistics
    if additional_stats:
        doc.add_heading("6. Additional Statistics", level=1)
        for stat_name, stat_df in additional_stats.items():
            if stat_df is not None and not stat_df.empty:
                doc.add_heading(stat_name, level=2)
                _add_dataframe_table(doc, stat_df)
                doc.add_paragraph("")

    # Section 7: Key Figures
    if include_images and images_dir:
        doc.add_page_break()
        doc.add_heading("7. Key Figures", level=1)

        images_path = Path(images_dir)
        if images_path.exists():
            # Define key images to include
            key_images = [
                ('residual_distributions.png', 'Residual Distributions'),
                ('accuracy_by_landcover.png', 'Accuracy by Land Cover'),
                ('accuracy_by_elevation.png', 'Accuracy by Elevation'),
                ('terrain_slope.png', 'Terrain Slope Analysis'),
                ('coverage_and_voids.png', 'Coverage and Voids'),
                ('radar_geometry_analysis.png', 'Radar Geometry Analysis'),
                ('summary_dashboard.png', 'Summary Dashboard'),
            ]

            for img_file, caption in key_images:
                img_path = images_path / img_file
                if img_path.exists():
                    doc.add_heading(caption, level=2)
                    try:
                        doc.add_picture(str(img_path), width=Inches(6.0))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        doc.add_paragraph(f"[Error loading image: {e}]")
                    doc.add_paragraph("")

    # Save document
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] Word document saved to: {output_path}")


def _add_dataframe_table(
    doc: Document,
    df: pd.DataFrame,
    index_label: Optional[str] = None
) -> None:
    """
    Add a pandas DataFrame as a formatted table to the document.

    Parameters
    ----------
    doc : Document
        The Word document object
    df : pd.DataFrame
        DataFrame to add as a table
    index_label : str, optional
        Label for the index column
    """
    # Reset index to include it as a column if it has a name or label provided
    df_display = df.copy()
    if index_label or df.index.name:
        df_display = df_display.reset_index()
        if index_label and 'index' in df_display.columns:
            df_display = df_display.rename(columns={'index': index_label})

    # Create table
    table = doc.add_table(rows=1, cols=len(df_display.columns))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(df_display.columns):
        header_cells[i].text = str(col_name)
        _format_cell(header_cells[i], bold=True)

    # Data rows
    for _, row in df_display.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            # Format numbers appropriately
            if isinstance(value, (int, np.integer)):
                row_cells[i].text = f"{value:,}"
            elif isinstance(value, (float, np.floating)):
                if abs(value) < 0.01 and value != 0:
                    row_cells[i].text = f"{value:.2e}"
                else:
                    row_cells[i].text = f"{value:.2f}"
            else:
                row_cells[i].text = str(value)


def _format_cell(cell, bold: bool = False, font_size: int = 10) -> None:
    """Apply formatting to a table cell."""
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = bold
    run.font.size = Pt(font_size)


def create_summary_statistics_dict(
    residuals_tin: np.ndarray,
    residuals_cop: np.ndarray,
    nmad_tin: float,
    nmad_cop: float
) -> Dict[str, Dict[str, float]]:
    """
    Create overall statistics dictionary from residual arrays.

    Parameters
    ----------
    residuals_tin : np.ndarray
        Residuals for TINItaly comparison
    residuals_cop : np.ndarray
        Residuals for Copernicus comparison
    nmad_tin : float
        NMAD for TINItaly
    nmad_cop : float
        NMAD for Copernicus

    Returns
    -------
    dict
        Dictionary with overall statistics for both reference DEMs
    """
    return {
        'TINItaly': {
            'n_points': len(residuals_tin),
            'bias': float(np.mean(residuals_tin)),
            'rmse': float(np.sqrt(np.mean(residuals_tin**2))),
            'nmad': float(nmad_tin),
            'min': float(np.min(residuals_tin)),
            'max': float(np.max(residuals_tin))
        },
        'Copernicus': {
            'n_points': len(residuals_cop),
            'bias': float(np.mean(residuals_cop)),
            'rmse': float(np.sqrt(np.mean(residuals_cop**2))),
            'nmad': float(nmad_cop),
            'min': float(np.min(residuals_cop)),
            'max': float(np.max(residuals_cop))
        }
    }
