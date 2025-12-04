"""
Insert appropriate figures into the Word document with captions.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

def insert_figures_into_document(doc_path, images_dir, output_path):
    """Insert figures into the Word document."""

    doc = Document(doc_path)
    images_dir = Path(images_dir)

    # Define figure mappings based on document content
    figure_mappings = [
        {
            'after_text': 'Figure 1: Spatial coverage of the SAOCOM',
            'image': 'coverage.png',
            'caption': 'Figure 1: Spatial coverage of the SAOCOM-1A and 1B StripMap acquisitions used in this study, overlaid on a relief map of Italy.',
            'width': 6.0
        },
        {
            'section': 'Results',
            'after_text': 'Overall Validation Statistics',
            'image': 'bland_altman.png',
            'caption': 'Figure 2: Bland-Altman plot showing agreement between SAOCOM and TINItaly elevations. The systematic bias and 95% limits of agreement indicate overall accuracy.',
            'width': 6.0
        },
        {
            'after_text': 'Accuracy by Slope Category',
            'image': 'accuracy_by_elevation.png',
            'caption': 'Figure 3: Accuracy metrics stratified by slope category, showing degradation in steep terrain.',
            'width': 5.5
        },
        {
            'after_text': 'Accuracy by Land Cover Type',
            'image': 'accuracy_by_landcover.png',
            'caption': 'Figure 4: Accuracy assessment across different land cover types, highlighting vegetation effects on L-band penetration.',
            'width': 6.0
        },
        {
            'after_text': 'Radar Geometry Statistics',
            'image': 'hexbin_density.png',
            'caption': 'Figure 5: Hexagonal binning density plot showing spatial distribution of elevation residuals.',
            'width': 5.5
        },
    ]

    print("Inserting figures into document...")
    print("=" * 70)

    inserted_count = 0

    for fig_map in figure_mappings:
        # Find the paragraph to insert after
        insert_index = None
        for i, para in enumerate(doc.paragraphs):
            if fig_map['after_text'] in para.text:
                insert_index = i + 1
                print(f"\nFound insertion point at paragraph {i}: {para.text[:60]}...")
                break

        if insert_index is None:
            print(f"[SKIP] Could not find insertion point for: {fig_map['after_text'][:50]}")
            continue

        # Check if image exists
        image_path = images_dir / fig_map['image']
        if not image_path.exists():
            print(f"[SKIP] Image not found: {fig_map['image']}")
            continue

        # Insert the image
        # We need to insert into the document at the correct position
        # This is tricky - we'll add after the target paragraph
        target_para = doc.paragraphs[insert_index - 1]

        # Add a new paragraph for the image
        new_para = target_para.insert_paragraph_before() if insert_index == 0 else None
        if new_para is None:
            # Insert after target
            # Get the element after target
            p_element = target_para._element
            new_p = p_element.addnext(target_para._element.__class__())
            new_para = doc.paragraphs[insert_index]

        # Actually, let's use a simpler approach - insert at end and we'll manually reorganize
        # For now, let's build a list of what to insert
        print(f"[OK] Would insert {fig_map['image']} with caption: {fig_map['caption'][:60]}...")
        inserted_count += 1

    print(f"\n{'=' * 70}")
    print(f"Identified {inserted_count} figures to insert")
    print("\nNote: Inserting images into existing Word docs requires manual positioning")
    print("Creating a new document with figures...")

    # Create new document with figures
    create_document_with_figures(doc, images_dir, figure_mappings, output_path)

def create_document_with_figures(original_doc, images_dir, figure_mappings, output_path):
    """Create a new document with figures inserted."""

    new_doc = Document()

    # Copy styles from original
    for style in original_doc.styles:
        try:
            if style.name not in new_doc.styles:
                pass  # Would need to copy style definitions
        except:
            pass

    # Track which figures we've inserted
    figures_inserted = set()

    # Go through original paragraphs
    for para in original_doc.paragraphs:
        # Add the paragraph
        new_para = new_doc.add_paragraph(para.text, style=para.style.name if para.style else None)
        try:
            new_para.alignment = para.alignment
        except (ValueError, AttributeError):
            pass  # Skip if alignment can't be copied

        # Check if we should insert a figure after this paragraph
        for fig_map in figure_mappings:
            if fig_map['image'] in figures_inserted:
                continue

            if fig_map['after_text'] in para.text:
                # Insert blank line
                new_doc.add_paragraph()

                # Insert image
                image_path = images_dir / fig_map['image']
                if image_path.exists():
                    new_doc.add_picture(str(image_path), width=Inches(fig_map['width']))

                    # Add caption
                    caption_para = new_doc.add_paragraph()
                    caption_para.add_run(fig_map['caption']).italic = True
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add blank line after
                    new_doc.add_paragraph()

                    figures_inserted.add(fig_map['image'])
                    print(f"[OK] Inserted {fig_map['image']}")

    # Copy tables
    for table in original_doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text

    # Save
    new_doc.save(output_path)
    print(f"\n[DONE] Saved document with figures to: {output_path}")

if __name__ == '__main__':
    doc_path = r'C:\users\colto\documents\github\saocom_project\Untitled 12.docx'
    images_dir = r'C:\users\colto\documents\github\saocom_project\images'
    output_path = r'C:\users\colto\documents\github\saocom_project\Untitled 12_with_figures.docx'

    insert_figures_into_document(doc_path, images_dir, output_path)
