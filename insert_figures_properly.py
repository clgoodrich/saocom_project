"""
Insert ALL figures with proper placement throughout the document.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def insert_figure(doc, image_path, caption, width=6.0):
    """Helper to insert a figure with caption."""
    doc.add_paragraph()  # Blank line

    # Insert image
    img_para = doc.add_paragraph()
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add caption
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run(caption)
    caption_run.italic = True
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Blank line after

def create_complete_document(doc_path, images_dir, output_path):
    """Create complete document with all figures."""

    doc = Document(doc_path)
    new_doc = Document()
    images_dir = Path(images_dir)

    fig_num = 1
    inserted = 0

    print("Creating comprehensive document with all figures...")
    print("=" * 70)

    for para_idx, para in enumerate(doc.paragraphs):
        # Copy paragraph
        new_para = new_doc.add_paragraph(para.text)
        try:
            if para.style:
                new_para.style = para.style.name
        except:
            pass

        # Insert figures at specific locations

        # After Figure 1 caption (para 34)
        if para_idx == 34:
            imgs = [
                ('spatial_coverage.png', 6.0, 'coverage.png'),
            ]
            for img, width, fallback in imgs:
                path = images_dir / img
                if not path.exists():
                    path = images_dir / fallback
                if path.exists():
                    insert_figure(new_doc, path, f'Figure {fig_num}: Spatial coverage of SAOCOM acquisitions over Italy.', width)
                    print(f"[{fig_num}] {img}")
                    fig_num += 1
                    inserted += 1

            # Add SAOCOM points map
            path = images_dir / 'saocom_points_map.png'
            if path.exists():
                insert_figure(new_doc, path, f'Figure {fig_num}: Distribution of SAOCOM point measurements.', 6.0)
                print(f"[{fig_num}] saocom_points_map.png")
                fig_num += 1
                inserted += 1

        # After Ancillary section (around para 41-43)
        if para_idx == 43:
            imgs = [
                ('land_cover_map.png', 6.0),
                ('land_cover_histograms.png', 6.0),
            ]
            captions = [
                'CORINE Land Cover classification across the study area.',
                'Distribution of SAOCOM points across land cover categories.'
            ]
            for (img, width), cap in zip(imgs, captions):
                path = images_dir / img
                if path.exists():
                    insert_figure(new_doc, path, f'Figure {fig_num}: {cap}', width)
                    print(f"[{fig_num}] {img}")
                    fig_num += 1
                    inserted += 1

        # After Table 1 (Overall Accuracy - para 72)
        if para_idx == 73:
            imgs = [
                ('bland_altman.png', 6.0, 'Bland-Altman plot showing SAOCOM vs TINItaly agreement.'),
                ('residual_distributions.png', 6.0, 'Distribution of elevation residuals.'),
                ('scatter_comparisons.png', 6.0, 'Scatter plots of SAOCOM vs reference DEMs.'),
                ('hexbin_density.png', 5.5, 'Hexagonal density plot of residuals.'),
                ('spatial_residuals.png', 6.0, 'Spatial distribution of elevation residuals.'),
                ('hist2d_comparison.png', 6.0, '2D histogram comparison of elevations.'),
            ]
            for img, width, cap in imgs:
                path = images_dir / img
                if path.exists():
                    insert_figure(new_doc, path, f'Figure {fig_num}: {cap}', width)
                    print(f"[{fig_num}] {img}")
                    fig_num += 1
                    inserted += 1

        # After Table 2 (Slope - para 77)
        if para_idx == 78:
            imgs = [
                ('terrain_slope.png', 5.5, 'Slope map showing terrain complexity.'),
                ('accuracy_by_elevation.png', 6.0, 'Accuracy metrics by slope category.'),
                ('violin_plot_slope.png', 6.0, 'Violin plots of residuals across slope categories.'),
            ]
            for img, width, cap in imgs:
                path = images_dir / img
                if path.exists():
                    insert_figure(new_doc, path, f'Figure {fig_num}: {cap}', width)
                    print(f"[{fig_num}] {img}")
                    fig_num += 1
                    inserted += 1

        # After Table 3 (Land Cover - para 82)
        if para_idx == 83:
            imgs = [
                ('accuracy_by_landcover.png', 6.0, 'Accuracy by major land cover types.'),
                ('accuracy_by_detailed_land_cover.png', 6.0, 'Detailed land cover accuracy metrics.'),
                ('land_cover_vs_terrain.png', 6.0, 'Interaction between land cover and terrain slope.'),
                ('land_cover_Broad-leaved_forest.png', 5.5, 'Residuals for Broad-leaved Forest.'),
                ('land_cover_Vineyards.png', 5.5, 'Residuals for Vineyards.'),
                ('land_cover_Pastures.png', 5.5, 'Residuals for Pastures.'),
                ('land_cover_Olive_groves.png', 5.5, 'Residuals for Olive Groves.'),
                ('land_cover_Complex_cultivation_patterns.png', 5.5, 'Residuals for Complex Cultivation.'),
            ]
            for img, width, cap in imgs:
                path = images_dir / img
                if path.exists():
                    insert_figure(new_doc, path, f'Figure {fig_num}: {cap}', width)
                    print(f"[{fig_num}] {img}")
                    fig_num += 1
                    inserted += 1

        # After Table 4 (Radar Geometry - para 89)
        if para_idx == 90:
            path = images_dir / 'radar_geometry_analysis.png'
            if path.exists():
                insert_figure(new_doc, path, f'Figure {fig_num}: Accuracy vs radar viewing geometry.', 6.0)
                print(f"[{fig_num}] radar_geometry_analysis.png")
                fig_num += 1
                inserted += 1

        # Discussion section (para 93+)
        if para_idx == 94:
            imgs = [
                ('reference_dem_comparison.png', 6.0, 'L-band vs X-band DEM comparison.'),
                ('gridded_comparison.png', 6.0, 'Gridded elevation differences.'),
                ('residuals_vs_coherence.png', 6.0, 'Coherence vs elevation accuracy.'),
                ('coverage_and_voids.png', 6.0, 'Data coverage and void analysis.'),
                ('summary_dashboard.png', 6.5, 'Summary dashboard of validation metrics.'),
            ]
            for img, width, cap in imgs:
                path = images_dir / img
                if path.exists():
                    insert_figure(new_doc, path, f'Figure {fig_num}: {cap}', width)
                    print(f"[{fig_num}] {img}")
                    fig_num += 1
                    inserted += 1

    # Copy tables
    print(f"\nCopying {len(doc.tables)} tables...")
    for table_idx, table in enumerate(doc.tables):
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        try:
            new_table.style = 'Light Grid Accent 1'
        except:
            pass

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text

    # Save
    new_doc.save(output_path)
    print(f"\n{'=' * 70}")
    print(f"[SUCCESS] Inserted {inserted} figures into document")
    print(f"[SUCCESS] Saved to: {output_path}")

if __name__ == '__main__':
    doc_path = r'C:\users\colto\documents\github\saocom_project\Untitled 12.docx'
    images_dir = r'C:\users\colto\documents\github\saocom_project\images'
    output_path = r'C:\users\colto\documents\github\saocom_project\Untitled 12_complete.docx'

    create_complete_document(doc_path, images_dir, output_path)
