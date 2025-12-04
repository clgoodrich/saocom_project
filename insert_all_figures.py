"""
Insert ALL relevant figures into the Word document with appropriate captions.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def insert_all_figures(doc_path, images_dir, output_path):
    """Insert comprehensive set of figures into document."""

    doc = Document(doc_path)
    images_dir = Path(images_dir)

    # Comprehensive figure mappings
    figure_mappings = [
        # Section 3.1 - SAOCOM Dataset
        {
            'after_text': 'Figure 1: Spatial coverage of the SAOCOM',
            'image': 'spatial_coverage.png',
            'caption': 'Figure 1: Spatial coverage of the SAOCOM-1A and 1B StripMap acquisitions over Italy.',
            'width': 6.0,
            'fallback': 'coverage.png'
        },
        {
            'after_text': 'Figure 1: Spatial coverage of the SAOCOM',
            'image': 'saocom_points_map.png',
            'caption': 'Figure 2: Distribution of SAOCOM point measurements across the study area.',
            'width': 6.0
        },

        # Section 3.3 - Ancillary Data
        {
            'after_text': '3.3 Ancillary Geospatial Layers',
            'image': 'land_cover_map.png',
            'caption': 'Figure 3: CORINE Land Cover classification across the study area, showing the diversity of surface types analyzed.',
            'width': 6.0
        },
        {
            'after_text': 'land_cover_map.png',
            'image': 'land_cover_histograms.png',
            'caption': 'Figure 4: Distribution of SAOCOM points across different land cover categories.',
            'width': 6.0
        },

        # Section 5.1 - Overall Accuracy
        {
            'after_text': 'Table 1: Overall Validation Statistics',
            'image': 'bland_altman.png',
            'caption': 'Figure 5: Bland-Altman plot showing agreement between SAOCOM and TINItaly elevations.',
            'width': 6.0
        },
        {
            'after_text': 'bland_altman.png',
            'image': 'residual_distributions.png',
            'caption': 'Figure 6: Distribution of elevation residuals for both TINItaly and Copernicus comparisons.',
            'width': 6.0
        },
        {
            'after_text': 'residual_distributions.png',
            'image': 'scatter_comparisons.png',
            'caption': 'Figure 7: Scatter plots comparing SAOCOM elevations against reference DEMs.',
            'width': 6.0
        },
        {
            'after_text': 'scatter_comparisons.png',
            'image': 'hexbin_density.png',
            'caption': 'Figure 8: Hexagonal density plot showing spatial distribution of elevation residuals.',
            'width': 5.5
        },
        {
            'after_text': 'hexbin_density.png',
            'image': 'spatial_residuals.png',
            'caption': 'Figure 9: Spatial distribution of elevation residuals across the study area.',
            'width': 6.0
        },

        # Section 5.2 - Impact of Slope
        {
            'after_text': 'Table 2: Accuracy by Slope Category',
            'image': 'terrain_slope.png',
            'caption': 'Figure 10: Slope map of the study area showing terrain complexity.',
            'width': 5.5
        },
        {
            'after_text': 'terrain_slope.png',
            'image': 'accuracy_by_elevation.png',
            'caption': 'Figure 11: Accuracy metrics stratified by slope category.',
            'width': 6.0
        },
        {
            'after_text': 'accuracy_by_elevation.png',
            'image': 'violin_plot_slope.png',
            'caption': 'Figure 12: Violin plots showing residual distributions across slope categories.',
            'width': 6.0
        },

        # Section 5.3 - Vegetation Penetration
        {
            'after_text': 'Table 3: Accuracy by Land Cover Type',
            'image': 'accuracy_by_landcover.png',
            'caption': 'Figure 13: Accuracy assessment across major land cover types.',
            'width': 6.0
        },
        {
            'after_text': 'accuracy_by_landcover.png',
            'image': 'accuracy_by_detailed_land_cover.png',
            'caption': 'Figure 14: Detailed accuracy metrics for specific land cover categories.',
            'width': 6.0
        },
        {
            'after_text': 'accuracy_by_detailed_land_cover.png',
            'image': 'land_cover_vs_terrain.png',
            'caption': 'Figure 15: Interaction between land cover type and terrain slope on accuracy.',
            'width': 6.0
        },

        # Individual land cover types
        {
            'after_text': 'land_cover_vs_terrain.png',
            'image': 'land_cover_Broad-leaved_forest.png',
            'caption': 'Figure 16: Residual analysis for Broad-leaved Forest, showing L-band penetration characteristics.',
            'width': 5.5
        },
        {
            'after_text': 'land_cover_Broad-leaved_forest.png',
            'image': 'land_cover_Vineyards.png',
            'caption': 'Figure 17: Residual analysis for Vineyards, demonstrating agricultural surface interactions.',
            'width': 5.5
        },
        {
            'after_text': 'land_cover_Vineyards.png',
            'image': 'land_cover_Pastures.png',
            'caption': 'Figure 18: Residual analysis for Pastures.',
            'width': 5.5
        },
        {
            'after_text': 'land_cover_Pastures.png',
            'image': 'land_cover_Olive_groves.png',
            'caption': 'Figure 19: Residual analysis for Olive Groves.',
            'width': 5.5
        },
        {
            'after_text': 'land_cover_Olive_groves.png',
            'image': 'land_cover_Complex_cultivation_patterns.png',
            'caption': 'Figure 20: Residual analysis for Complex Cultivation Patterns.',
            'width': 5.5
        },

        # Section 5.4 - Radar Geometry
        {
            'after_text': 'Table 4: Radar Geometry Statistics',
            'image': 'radar_geometry_analysis.png',
            'caption': 'Figure 21: Analysis of accuracy as a function of radar viewing geometry.',
            'width': 6.0
        },

        # Section 6 - Discussion
        {
            'after_text': '6.1 L-Band vs. X-Band Synergy',
            'image': 'reference_dem_comparison.png',
            'caption': 'Figure 22: Comparison of SAOCOM (L-band) against Copernicus (X-band) reference.',
            'width': 6.0
        },
        {
            'after_text': 'reference_dem_comparison.png',
            'image': 'gridded_comparison.png',
            'caption': 'Figure 23: Gridded elevation comparison showing systematic differences between sensors.',
            'width': 6.0
        },

        # Additional analysis
        {
            'after_text': 'gridded_comparison.png',
            'image': 'residuals_vs_coherence.png',
            'caption': 'Figure 24: Relationship between InSAR coherence and elevation accuracy.',
            'width': 6.0
        },
        {
            'after_text': 'residuals_vs_coherence.png',
            'image': 'coverage_and_voids.png',
            'caption': 'Figure 25: Coverage analysis showing data voids due to low coherence.',
            'width': 6.0
        },
        {
            'after_text': 'coverage_and_voids.png',
            'image': 'summary_dashboard.png',
            'caption': 'Figure 26: Summary dashboard of key validation metrics.',
            'width': 6.5
        },
    ]

    # Create new document
    new_doc = Document()
    figures_inserted = set()
    figure_counter = 1

    print(f"Inserting figures into document...")
    print("=" * 70)

    # Go through original paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        # Add the paragraph
        new_para = new_doc.add_paragraph(para.text)
        try:
            new_para.alignment = para.alignment
            if para.style:
                new_para.style = para.style.name
        except (ValueError, AttributeError):
            pass

        # Check if we should insert figure(s) after this paragraph
        for fig_map in figure_mappings:
            if fig_map['image'] in figures_inserted:
                continue

            if fig_map['after_text'] in para.text:
                # Try main image first, fallback if specified
                image_path = images_dir / fig_map['image']
                if not image_path.exists() and 'fallback' in fig_map:
                    image_path = images_dir / fig_map['fallback']

                if image_path.exists():
                    # Insert blank line
                    new_doc.add_paragraph()

                    # Insert image
                    img_para = new_doc.add_paragraph()
                    run = img_para.add_run()
                    run.add_picture(str(image_path), width=Inches(fig_map['width']))
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add caption
                    caption_para = new_doc.add_paragraph()
                    caption_run = caption_para.add_run(fig_map['caption'])
                    caption_run.italic = True
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add blank line after
                    new_doc.add_paragraph()

                    figures_inserted.add(fig_map['image'])
                    print(f"[{figure_counter}] Inserted {fig_map['image']}")
                    figure_counter += 1
                else:
                    print(f"[SKIP] Image not found: {fig_map['image']}")

    # Copy tables
    print(f"\nCopying {len(doc.tables)} tables...")
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = 'Light Grid Accent 1'

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text

    # Save
    new_doc.save(output_path)
    print(f"\n{'=' * 70}")
    print(f"[DONE] Inserted {len(figures_inserted)} figures")
    print(f"[DONE] Saved to: {output_path}")

if __name__ == '__main__':
    doc_path = r'C:\users\colto\documents\github\saocom_project\Untitled 12.docx'
    images_dir = r'C:\users\colto\documents\github\saocom_project\images'
    output_path = r'C:\users\colto\documents\github\saocom_project\Untitled 12_complete.docx'

    insert_all_figures(doc_path, images_dir, output_path)
